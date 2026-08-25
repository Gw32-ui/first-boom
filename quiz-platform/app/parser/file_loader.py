"""文件文本抽取：docx / pdf / txt / md → 纯文本。"""
from __future__ import annotations

import io
import re
import uuid
import zipfile
from pathlib import Path

from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
V_NS = "urn:schemas-microsoft-com:vml"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
DECORATIVE_IMG_THRESHOLD = 5  # 同一图片出现在 ≥5 个段落 → 视为水印/装饰，跳过
# 浏览器可直接渲染的图片格式；EMF/WMF/TIFF 等需转换后再落盘
DOCX_RENDERABLE_EXT = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".svg"}

# 待办③：相邻上标/下标片段（如 $^{-}$$^{11}$）合并成单个 $^{-11}$
SCRIPT_SUPER = re.compile(r"^\$\^\{([^}]*)\}\$$")
SCRIPT_SUB = re.compile(r"^\$\_\{([^}]*)\}\$$")

# OMML m:func 的常见函数名 → LaTeX 命令
FUNC_NAMES = {
    "sin": r"\sin", "cos": r"\cos", "tan": r"\tan", "cot": r"\cot",
    "sec": r"\sec", "csc": r"\csc", "log": r"\log", "ln": r"\ln",
    "exp": r"\exp", "lim": r"\lim", "lg": r"\lg",
}


def extract_bytes_text(data: bytes, suffix: str) -> str:
    """从内存字节抽取文本（网页上传用）。"""
    suffix = suffix.lower()
    if suffix in (".docx", ".doc"):
        from docx import Document

        doc = Document(io.BytesIO(data))
        lines: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return "\n".join(lines)
    if suffix == ".pdf":
        import pdfplumber

        pages: list[str] = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                pages.append(page.extract_text() or "")
        return "\n\n".join(pages)
    if suffix in (".txt", ".md", ".text"):
        for enc in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                return data.decode(enc)
            except UnicodeDecodeError:
                continue
        return data.decode("utf-8", errors="replace")
    raise ValueError(f"不支持的文件类型: {suffix}")


def extract_file_text(path: str | Path) -> str:
    """按扩展名读取文件并抽取文本。"""
    p = Path(path)
    return extract_bytes_text(p.read_bytes(), p.suffix)


def extract_images(data: bytes, suffix: str, out_dir: str | Path) -> list[str]:
    """抽取 docx/pdf 中的内嵌图片，保存到 out_dir，返回图片文件路径列表。

    - docx：直接读取 word/media/ 下的图片原始字节
    - pdf：从页面图片流提取（支持 JPEG/DCTDecode；FlateDecode 暂跳过）
    - 其他类型：返回空列表
    """
    out = Path(out_dir)
    out.mkdir(parents=True, exist_ok=True)
    suffix = suffix.lower()
    prefix = f"img_{uuid.uuid4().hex[:8]}"
    if suffix == ".docx":
        return _extract_docx_images(data, out, prefix)
    if suffix == ".pdf":
        return _extract_pdf_images(data, out, prefix)
    return []


def _extract_docx_images(data: bytes, out_dir: Path, prefix: str) -> list[str]:
    """从 docx（zip）里取出 word/media/ 下的图片并落盘。"""
    saved: list[str] = []
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        media = sorted(n for n in z.namelist() if n.startswith("word/media/"))
        for i, name in enumerate(media):
            ext = Path(name).suffix.lower() or ".bin"
            fname = f"{prefix}_{i:04d}{ext}"
            (out_dir / fname).write_bytes(z.read(name))
            saved.append(str(out_dir / fname))
    return saved


def _extract_pdf_images(data: bytes, out_dir: Path, prefix: str) -> list[str]:
    """从 PDF 页面图片流提取图片（JPEG 直接落盘，其他编码跳过）。"""
    import pdfplumber
    from PIL import Image

    saved: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        n = 0
        for page in pdf.pages:
            for img in page.images:
                stream = img.get("stream")
                raw = getattr(stream, "rawdata", None) if stream is not None else None
                if not raw:
                    continue
                try:
                    pil_img = Image.open(io.BytesIO(raw))
                    pil_img.load()
                except Exception:
                    continue
                fname = f"{prefix}_{n:04d}.png"
                pil_img.save(out_dir / fname)
                saved.append(str(out_dir / fname))
                n += 1
    return saved


def _omml_to_latex(el) -> str:
    """把单个 OMML 数学元素转成 LaTeX（基础版：分式/上下标/根式/求和/括号）。"""
    tag = etree.QName(el).localname

    def find(child: str):
        return el.find(f"{{{M_NS}}}{child}")

    def base(s: str) -> str:
        """底数包装：单字符不加括号，多字符/命令加括号。"""
        if not s:
            return "{}"
        if len(s) == 1 and not s.startswith("\\"):
            return s
        return "{" + s + "}"

    if tag in ("oMath", "oMathPara"):
        # 清理 Word 插入的不可见字符（U+2061 函数应用符等）
        return "".join(_omml_to_latex(c) for c in el).replace("\u2061", "")
    if tag == "f":  # 分式
        num, den = find("num"), find("den")
        return (
            "\\frac{"
            + (_omml_to_latex(num) if num is not None else "")
            + "}{"
            + (_omml_to_latex(den) if den is not None else "")
            + "}"
        )
    if tag == "func":  # 函数：fName 后跟括号参数
        fname = find("fName")
        e = find("e")
        fn = (_omml_to_latex(fname) if fname is not None else "").strip()
        fn = FUNC_NAMES.get(fn, fn + " ")
        return fn + "(" + (_omml_to_latex(e) if e is not None else "") + ")"
    if tag == "sSup":  # 上标
        e, sup = find("e"), find("sup")
        return (
            base(_omml_to_latex(e) if e is not None else "")
            + "^{"
            + (_omml_to_latex(sup) if sup is not None else "")
            + "}"
        )
    if tag == "sSub":  # 下标
        e, sub = find("e"), find("sub")
        return (
            base(_omml_to_latex(e) if e is not None else "")
            + "_{"
            + (_omml_to_latex(sub) if sub is not None else "")
            + "}"
        )
    if tag == "sSubSup":  # 上下标
        e, sub, sup = find("e"), find("sub"), find("sup")
        return (
            base(_omml_to_latex(e) if e is not None else "")
            + "_{"
            + (_omml_to_latex(sub) if sub is not None else "")
            + "}^{"
            + (_omml_to_latex(sup) if sup is not None else "")
            + "}"
        )
    if tag == "rad":  # 根式
        e, deg = find("e"), find("deg")
        inner = _omml_to_latex(e) if e is not None else ""
        if deg is not None and len(list(deg)):
            return "\\sqrt[" + _omml_to_latex(deg) + "]{" + inner + "}"
        return "\\sqrt{" + inner + "}"
    if tag == "nary":  # 求和/积分
        symbol = "\\sum"
        nary_pr = find("naryPr")
        if nary_pr is not None:
            chr_el = nary_pr.find(f"{{{M_NS}}}chr")
            val = chr_el.get(f"{{{M_NS}}}val") if chr_el is not None else None
            if val == "∫":
                symbol = "\\int"
            elif val == "∏":
                symbol = "\\prod"
        sub, sup, e = find("sub"), find("sup"), find("e")
        sub_l = _omml_to_latex(sub) if sub is not None else ""
        sup_l = _omml_to_latex(sup) if sup is not None else ""
        e_l = _omml_to_latex(e) if e is not None else ""
        if sub_l and sup_l:
            return f"{symbol}_{{{sub_l}}}^{{{sup_l}}} {e_l}"
        if sub_l:
            return f"{symbol}_{{{sub_l}}} {e_l}"
        if sup_l:
            return f"{symbol}^{{{sup_l}}} {e_l}"
        return f"{symbol} {e_l}"
    if tag == "d":  # 分隔符（括号，支持自定义 begChr/endChr）
        dpr = find("dPr")
        beg = end = None
        if dpr is not None:
            b = dpr.find(f"{{{M_NS}}}begChr")
            en = dpr.find(f"{{{M_NS}}}endChr")
            beg = b.get(f"{{{M_NS}}}val") if b is not None else None
            end = en.get(f"{{{M_NS}}}val") if en is not None else None
        inner = "".join(_omml_to_latex(c) for c in el)
        return (beg or "(") + inner + (end or ")")
    if tag == "r":
        t = find("t")
        return (t.text or "") if t is not None else ""
    if tag == "t":
        return el.text or ""
    # 其余标签（m:mr、m:egArr 等）：递归拼接子元素
    return "".join(_omml_to_latex(c) for c in el)


def extract_math_latex(data: bytes) -> list[str]:
    """从 docx 提取所有公式（OMML）并转成 LaTeX 字符串，按文档顺序返回。"""
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        xml_bytes = z.read("word/document.xml")
    root = etree.fromstring(xml_bytes)

    results: list[str] = []
    for p in root.iter(f"{{{W_NS}}}p"):
        for el in p.iter():
            ln = etree.QName(el).localname
            if ln == "oMathPara":
                for om in el:
                    if etree.QName(om).localname == "oMath":
                        results.append(_omml_to_latex(om))
                continue
            if ln == "oMath":
                parent = el.getparent()
                pn = etree.QName(parent).localname if parent is not None else ""
                if pn in ("oMath", "oMathPara"):
                    continue
                results.append(_omml_to_latex(el))
    return results


def _image_r_id(el) -> str | None:
    """从 pict（VML imagedata）或 drawing（blip）取图片关系 id。"""
    for im in el.iter(f"{{{V_NS}}}imagedata"):
        rid = im.get(f"{{{R_NS}}}id")
        if rid:
            return rid
    for blip in el.iter(f"{{{A_NS}}}blip"):
        rid = blip.get(f"{{{R_NS}}}embed")
        if rid:
            return rid
    return None


def _count_para_images(root, relmap: dict[str, str]) -> dict[str, int]:
    """统计每个图片目标出现在多少个段落（用于过滤水印/装饰图）。"""
    counts: dict[str, int] = {}
    for p in root.iter(f"{{{W_NS}}}p"):
        seen: set[str] = set()
        for el in p.iter():
            rid = _image_r_id(el)
            if rid:
                target = relmap.get(rid)
                if target:
                    seen.add(Path(target).name)
        for name in seen:
            counts[name] = counts.get(name, 0) + 1
    return counts


def _merge_adjacent_scripts(segs: list[str]) -> list[str]:
    """合并相邻的上标/下标片段。

    docx 里同一个上标可能被拆成多个 run（如 “－” 和 “11” 分开），
    各自转成 $^{...}$ 后应合并为 $^{-11}$，避免渲染成两段错位的上标。
    """
    out: list[str] = []
    for s in segs:
        ms = SCRIPT_SUPER.match(s)
        if ms and out:
            prev = SCRIPT_SUPER.match(out[-1])
            if prev:
                out[-1] = "$^{" + prev.group(1) + ms.group(1) + "}$"
                continue
        mb = SCRIPT_SUB.match(s)
        if mb and out:
            prev = SCRIPT_SUB.match(out[-1])
            if prev:
                out[-1] = "$_{" + prev.group(1) + mb.group(1) + "}$"
                continue
        out.append(s)
    return out

def _convert_docx_media(
    data: bytes,
    ext: str,
    out_dir: Path,
    fname: str,
) -> str | None:
    """把 docx 媒体写成可被浏览器渲染的文件。

    - png/jpg/jpeg/gif/webp/bmp/svg 直接落盘
    - EMF/WMF/TIFF 等转 PNG（Windows 下 Pillow 走 GDI）
    - 转换失败时保留原始文件并返回 None（该图需人工处理）
    """
    ext = (ext or "").lower()
    if ext in DOCX_RENDERABLE_EXT:
        (out_dir / fname).write_bytes(data)
        return fname
    try:
        from PIL import Image

        im = Image.open(io.BytesIO(data))
        im.load()
        if im.mode not in ("RGB", "RGBA"):
            im = im.convert("RGB")
        png_name = str(Path(fname).with_suffix(".png"))
        im.save(out_dir / png_name)
        return png_name
    except Exception:  # noqa: BLE001 - 单图失败保留原文件，不中断
        (out_dir / fname).write_bytes(data)
        return None


def _save_media(
    name: str,
    media: dict[str, bytes],
    out_dir: Path,
    prefix: str,
    unrendered: list[str],
) -> str:
    """保存 docx 媒体并返回锚点文件名；无法渲染的文件名记入 unrendered。"""
    data = media.get(name)
    if data is None:
        return ""
    fname = f"{prefix}_{name}"
    ext = Path(name).suffix.lower()
    if ext not in DOCX_RENDERABLE_EXT:
        png_variant = str(Path(fname).with_suffix(".png"))
        if (out_dir / png_variant).exists():
            return png_variant
    if (out_dir / fname).exists():
        return fname
    written = _convert_docx_media(data, ext, out_dir, fname)
    if written is None:
        unrendered.append(fname)
        return fname
    return written


def _run_segments(
    run_el,
    relmap: dict[str, str],
    media: dict[str, bytes],
    out_dir: Path,
    para_img_counts: dict[str, int],
    prefix: str,
    unrendered: list[str],
) -> list[str]:
    """把单个 w:r 转成文本片段（含上下标/行内图片）。"""
    segs: list[str] = []
    va = run_el.find(f".//{{{W_NS}}}vertAlign")
    va_val = va.get(f"{{{W_NS}}}val") if va is not None else None
    for child in run_el:
        ln = etree.QName(child).localname
        if ln == "t":
            txt = child.text or ""
            if not txt:
                continue
            if va_val == "superscript":
                segs.append("$^{" + txt.replace("－", "-").replace("−", "-") + "}$")
            elif va_val == "subscript":
                segs.append("$_{" + txt + "}$")
            else:
                segs.append(txt)
        elif ln == "tab":
            segs.append(" ")
        elif ln in ("pict", "drawing"):
            rid = _image_r_id(child)
            if rid:
                target = relmap.get(rid)
                if target:
                    name = Path(target).name
                    if para_img_counts.get(name, 0) < DECORATIVE_IMG_THRESHOLD:
                        fname = _save_media(
                            name, media, out_dir, prefix, unrendered
                        )
                        if fname:
                            segs.append(f"【图:{fname}】")
    return segs


def _para_segments(
    p,
    relmap: dict[str, str],
    media: dict[str, bytes],
    out_dir: Path,
    para_img_counts: dict[str, int],
    prefix: str,
    unrendered: list[str],
) -> str:
    segs: list[str] = []
    for child in p:
        tag = child.tag
        ln = etree.QName(child).localname
        if tag == f"{{{W_NS}}}r":
            segs.extend(
                _run_segments(
                    child,
                    relmap,
                    media,
                    out_dir,
                    para_img_counts,
                    prefix,
                    unrendered,
                )
            )
        elif tag == f"{{{W_NS}}}hyperlink":
            for r_el in child.findall(f"{{{W_NS}}}r"):
                segs.extend(
                    _run_segments(
                        r_el,
                        relmap,
                        media,
                        out_dir,
                        para_img_counts,
                        prefix,
                        unrendered,
                    )
                )
        elif tag == f"{{{M_NS}}}oMath":
            segs.append("$" + _omml_to_latex(child) + "$")
        elif tag == f"{{{M_NS}}}oMathPara":
            for om in child.findall(f"{{{M_NS}}}oMath"):
                segs.append("$" + _omml_to_latex(om) + "$")
        elif ln in ("pict", "drawing"):
            rid = _image_r_id(child)
            if rid:
                target = relmap.get(rid)
                if target:
                    name = Path(target).name
                    if para_img_counts.get(name, 0) < DECORATIVE_IMG_THRESHOLD:
                        fname = _save_media(
                            name, media, out_dir, prefix, unrendered
                        )
                        if fname:
                            segs.append(f"【图:{fname}】")
        elif ln in (
            "pPr",
            "bookmarkStart",
            "bookmarkEnd",
            "commentRangeStart",
            "commentRangeEnd",
        ):
            continue
        else:
            # w:smartTag / w:ins 等：收集内部 run
            for r_el in child.iter(f"{{{W_NS}}}r"):
                segs.extend(
                    _run_segments(
                        r_el,
                        relmap,
                        media,
                        out_dir,
                        para_img_counts,
                        prefix,
                        unrendered,
                    )
                )
    return "".join(_merge_adjacent_scripts(segs)).strip()


def extract_docx_rich(
    data: bytes, images_dir: str | Path
) -> tuple[str, list[str]]:
    """按文档顺序把 docx 转成富文本。

    - 行内公式：$LaTeX$（OMML）
    - 图片：按出现位置输出 【图:文件名】并保存到 images_dir
      文件名带本次文档前缀（doc_xxxx_media），避免不同文档同名媒体互相覆盖；
      EMF/WMF/TIFF 尝试转 PNG，转换失败的图片名记入返回的 unrendered 列表
    - 上标/下标：$^{...}$ / $_{...}$
    - 同一图片出现在 ≥5 个段落视为水印/装饰，跳过
    """
    out = Path(images_dir)
    out.mkdir(parents=True, exist_ok=True)
    prefix = f"doc_{uuid.uuid4().hex[:8]}"
    unrendered: list[str] = []
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        xml_bytes = z.read("word/document.xml")
        rels_xml = z.read("word/_rels/document.xml.rels")
        media = {
            Path(n).name: z.read(n)
            for n in z.namelist()
            if n.startswith("word/media/") and not n.endswith("/")
        }
    root = etree.fromstring(xml_bytes)
    rels = etree.fromstring(rels_xml)
    relmap = {r.get("Id"): r.get("Target") for r in rels}
    para_img_counts = _count_para_images(root, relmap)

    lines: list[str] = []
    for p in root.iter(f"{{{W_NS}}}p"):
        line = _para_segments(
            p,
            relmap,
            media,
            out,
            para_img_counts,
            prefix,
            unrendered,
        )
        if line:
            lines.append(line)
    return "\n".join(lines), unrendered


# ---------------------------------------------------------------------------
# PDF 坐标级图题匹配（PyMuPDF）
# ---------------------------------------------------------------------------


def _save_pdf_image_bytes(data: bytes, ext: str, out_dir: Path, fname: str) -> bool:
    """把 PDF 图片字节转成 PNG 落盘（CMYK/带通道图统一转 RGB）。"""
    try:
        from PIL import Image

        im = Image.open(io.BytesIO(data))
        im.load()
        if im.mode in ("CMYK", "P", "LA", "RGBA", "PA"):
            im = im.convert("RGB")
        elif im.mode != "RGB":
            im = im.convert("RGB")
        im.save(out_dir / fname, "PNG")
        return True
    except Exception:  # noqa: BLE001 - 单图失败不中断
        return False


def _extract_pdf_bitmaps(page, doc, out_dir: Path, prefix: str, counter: list[int]) -> list[dict]:
    """提取 PDF 页面内嵌位图，返回 [{rect, fname}]。"""
    results: list[dict] = []
    try:
        # 新版 PyMuPDF 默认不返回 xref，必须显式开启；旧版兼容降级
        try:
            infos = page.get_image_info(xrefs=True)
        except Exception:  # noqa: BLE001 - 旧版 API 无 xrefs 参数
            infos = page.get_image_info()
    except Exception:  # noqa: BLE001 - API 差异兼容
        return results
    for info in infos:
        bbox = info.get("bbox")
        xref = info.get("xref", 0)
        if not bbox or not xref:
            continue
        try:
            ext_info = doc.extract_image(xref)
        except Exception:  # noqa: BLE001
            continue
        if not ext_info:
            continue
        raw = ext_info.get("image")
        if not raw:
            continue
        ext = (ext_info.get("ext") or "png").lower()
        if ext not in ("png", "jpg", "jpeg", "bmp", "tif", "tiff", "gif"):
            ext = "png"
        suffix = ext if ext == "png" else ext
        fname = f"{prefix}_{counter[0]:04d}_{ext}.png"
        counter[0] += 1
        if not _save_pdf_image_bytes(raw, ext, out_dir, fname):
            continue
        results.append({"rect": tuple(float(v) for v in bbox), "fname": fname})
    return results


def _extract_pdf_vectors(page, out_dir: Path, prefix: str, counter: list[int]) -> list[dict]:
    """渲染 PDF 页面中的矢量绘图（可无内嵌位图，如电路图/受力分析）。

    策略：收集所有绘图 path 的矩形，按 y 聚类过滤掉下划线/表格线等细小条带，
    再用 page.get_pixmap(clip=区域) 渲染成 PNG。
    """
    try:
        import pymupdf

        drawings = page.get_drawings()
    except Exception:  # noqa: BLE001
        return []
    if not drawings:
        return []
    rects = []
    for dr in drawings:
        r = dr.get("rect")
        if not r:
            continue
        w = r.x1 - r.x0
        h = r.y1 - r.y0
        if w <= 3 or h <= 3:  # 过滤点/细线
            continue
        # 过滤明显为下划线/表格线的窄条
        if (h < 4 and w > 20) or (w < 4 and h > 20):
            continue
        rects.append((r.x0, r.y0, r.x1, r.y1))
    if not rects:
        return []
    # y 聚类：把 y 距离相近的矩形并成一张"图区域"
    rects.sort(key=lambda r: (r[1], r[0]))
    clusters: list[list[tuple[float, float, float, float]]] = []
    cur: list[tuple[float, float, float, float]] = []
    last_y = None
    for r in rects:
        if last_y is None or abs(r[1] - last_y) < 25:
            cur.append(r)
        else:
            clusters.append(cur)
            cur = [r]
        last_y = r[1]
    if cur:
        clusters.append(cur)

    results: list[dict] = []
    for cl in clusters:
        x0 = min(r[0] for r in cl)
        y0 = min(r[1] for r in cl)
        x1 = max(r[2] for r in cl)
        y1 = max(r[3] for r in cl)
        w, h = x1 - x0, y1 - y0
        # 过滤太小的散点簇（可能是字符装饰）；保留宽度/高度达标的区域
        if w < 8 or h < 8:
            continue
        if w * h < 400:  # 过小面积不视为图
            continue
        try:
            clip = pymupdf.Rect(x0 - 1, y0 - 1, x1 + 1, y1 + 1)
            mat = pymupdf.Matrix(3, 3)  # 3x 缩放，清晰渲染
            pix = page.get_pixmap(matrix=mat, clip=clip)
            fname = f"{prefix}_{counter[0]:04d}_vec.png"
            counter[0] += 1
            if clip.width >= 2 and clip.height >= 2:
                pix.save(out_dir / fname)
                results.append({"rect": (x0, y0, x1, y1), "fname": fname})
        except Exception:  # noqa: BLE001 - 单区域失败不中断
            continue
    return results


def extract_pdf_pages(
    data: bytes,
    images_dir: str | Path,
    dpi: int = 200,
    max_pages: int = 10,
) -> list[str]:
    """把 PDF 每页整体渲染成 PNG，返回图片路径列表。

    公式型 PDF 的公式字形不是内嵌位图，extract_pdf_rich 提取不到图片；
    文本层碎片化时用整页渲染图作为视觉 OCR 输入。
    """
    import pymupdf

    out = Path(images_dir)
    out.mkdir(parents=True, exist_ok=True)
    prefix = f"page_{uuid.uuid4().hex[:6]}"
    saved: list[str] = []
    doc = pymupdf.open(stream=data, filetype="pdf")
    try:
        n = min(len(doc), max(1, max_pages))
        for i in range(n):
            page = doc.load_page(i)
            mat = pymupdf.Matrix(dpi / 72.0, dpi / 72.0)
            pix = page.get_pixmap(matrix=mat)
            fname = f"{prefix}_{i + 1:03d}.png"
            pix.save(out / fname)
            saved.append(str(out / fname))
    finally:
        doc.close()
    return saved


# ---------------------------------------------------------------------------
# PDF 图片归属映射（OCR 兜底时保留图片不丢）
# ---------------------------------------------------------------------------

# 与 model_parser.QUESTION_START 对齐的题号识别（取第一个数字作为归一化题号）
PDF_QUESTION_MARKER = re.compile(
    r"^\s*(?:第\s*(\d+)\s*题"
    r"|【\s*(\d{1,3})\s*[-－—]\s*\d{1,3}\s*】"
    r"|(\d{1,3})\s*[\.．、)）：:])"
)


def _pdf_qnum_key(text: str) -> str | None:
    """从文本块行首提取归一化题号（“1．”“第5题”“【4-4】” → “1”“5”“4”）。"""
    m = PDF_QUESTION_MARKER.match(text.strip())
    if not m:
        return None
    for group in m.groups():
        if group:
            return str(int(group))
    return None


def map_pdf_images_to_questions(
    image_events: list[dict],
    question_markers: list[dict],
) -> dict[str, str | None]:
    """把 PDF 图片按坐标归属到题号（读取顺序：页 → y → 题号优先）。

    - image_events: [{"fname": str, "page": int, "y": float, ...}]
    - question_markers: [{"key": str, "page": int, "y": float}]

    返回 {图片文件名: 归一化题号}；第一个题号之前的图片归属 None，
    由调用方按顺序兜底绑定并标待复核。
    """
    events: list[tuple[int, float, int, str]] = []
    for i, m in enumerate(question_markers):
        events.append((m["page"], m["y"], 0, m["key"]))
    for im in image_events:
        events.append((im["page"], im["y"], 1, im["fname"]))
    events.sort(key=lambda e: (e[0], e[1], e[2]))
    result: dict[str, str | None] = {}
    cur: str | None = None
    for _page, _y, kind, value in events:
        if kind == 0:
            cur = value
        else:
            result[value] = cur
    return result


def extract_pdf_rich_meta(
    data: bytes,
    images_dir: str | Path,
) -> tuple[str, list[str], list[dict], list[dict]]:
    """extract_pdf_rich 的元数据版：额外返回图片事件与题号标记。

    返回 (富文本, 图片路径列表, image_events, question_markers)：
    - image_events: [{"fname", "page", "y", "rect"}]
    - question_markers: [{"key", "page", "y"}]
    供 OCR 兜底路径在文本被替换后仍能按坐标把图片回填到对应题目。
    """
    import pymupdf

    out = Path(images_dir)
    out.mkdir(parents=True, exist_ok=True)
    prefix = f"img_{uuid.uuid4().hex[:8]}"
    counter: list[int] = [0]
    saved: list[str] = []
    lines: list[str] = []
    image_events: list[dict] = []
    question_markers: list[dict] = []

    doc = pymupdf.open(stream=data, filetype="pdf")
    try:
        for page_index in range(len(doc)):
            page = doc.load_page(page_index)
            page_w = page.rect.width
            page_h = page.rect.height
            # 收集图片事件（位图 + 矢量图），带 bbox
            img_events: list[dict] = []
            for bm in _extract_pdf_bitmaps(page, doc, out, prefix, counter):
                img_events.append(bm)
                saved.append(str(out / bm["fname"]))
            for vec in _extract_pdf_vectors(page, out, prefix, counter):
                img_events.append(vec)
                saved.append(str(out / vec["fname"]))
            # 文本块
            text_blocks: list[tuple[float, float, float, float, str]] = []
            try:
                blocks = page.get_text("blocks")  # (x0,y0,x1,y1,text,no,type)
                for b in blocks:
                    if len(b) < 5:
                        continue
                    x0, y0, x1, y1 = float(b[0]), float(b[1]), float(b[2]), float(b[3])
                    txt = (b[4] or "").strip()
                    if txt:
                        text_blocks.append((x0, y0, x1, y1, txt))
            except Exception:  # noqa: BLE001
                text_blocks = []

            # 合并排序：图片与文本块按 (top, left) 排序
            events: list[tuple[str, float, float, object]] = []
            for tb in text_blocks:
                events.append(("text", tb[1], tb[0], tb))
            for ie in img_events:
                events.append(("img", ie["rect"][1], ie["rect"][0], ie))
            events.sort(key=lambda e: (round(e[1], 1), e[2]))

            page_lines: list[str] = []
            for etype, _, _, payload in events:
                if etype == "text":
                    txt = payload[4]
                    key = _pdf_qnum_key(txt)
                    if key is not None:
                        question_markers.append(
                            {"key": key, "page": page_index, "y": payload[1]}
                        )
                    page_lines.append(txt)
                else:
                    # 装饰/水印过滤：图片特别小（<页面面积 0.3%）或贴近页面边缘整幅出现 → 视为装饰
                    r = payload["rect"]
                    iw = r[2] - r[0]
                    ih = r[3] - r[1]
                    if iw >= page_w * 0.85 and ih >= page_h * 0.85:
                        continue
                    image_events.append(
                        {
                            "fname": payload["fname"],
                            "page": page_index,
                            "y": payload["rect"][1],
                            "rect": payload["rect"],
                        }
                    )
                    page_lines.append(f"【图:{payload['fname']}】")
            if page_lines:
                lines.append("\n".join(page_lines))
                lines.append("")  # 页间空行，便于切块
    finally:
        doc.close()
    return "\n".join(lines).strip(), saved, image_events, question_markers


def extract_pdf_rich(data: bytes, images_dir: str | Path) -> tuple[str, list[str]]:
    """PDF → 富文本（图片按坐标位置插入【图:文件名】锚点）+ 图片路径列表。

    与 extract_docx_rich 对齐：图片锚点跟随正文，后续 model_parser 切题时
    图片自然归属到"它所在位置的那道题"，不再靠"按图片流顺序硬绑序号"。

    - 位图（内嵌 jpg/png 等）：直接提取落盘
    - 矢量绘图（电路图/受力分析等）：按绘图区域渲染成 PNG
    - 文本块与图片按 y 坐标排序，图片锚点插入到正确的文本位置
    """
    text, saved, _image_events, _markers = extract_pdf_rich_meta(data, images_dir)
    return text, saved
